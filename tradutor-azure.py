import requests
from docx import Document
import os

subscription_key = '6heKDjunL76YjCcrA3KM3NgD6SnFeObvyvnxx66betSHIEkneonHJQQJ99AKAC8vTInXJ3w3AAAbACOG2uUD'
endpoint = 'https://api.cognitive.microsofttranslator.com/'
location = 'westus2'
target_language = 'pt-BR'

def traduzir_documento(documento):
    doc = Document(documento)
    translated_doc = Document()
    
    # Construct the request headers
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }
    
    # API endpoint for translations
    path = '/translate'
    constructed_url = endpoint + path
    
    # Parameters for the request
    params = {
        'api-version': '3.0',
        'to': target_language
    }
    
    # Process each paragraph in the document
    for para in doc.paragraphs:
        if para.text.strip():  # Only translate non-empty paragraphs
            # Prepare the body
            body = [{
                'text': para.text
            }]
            
            # Make the API call
            try:
                response = requests.post(
                    constructed_url,
                    params=params,
                    headers=headers,
                    json=body
                )
                response.raise_for_status()
                
                # Get the translation
                translations = response.json()
                translated_text = translations[0]['translations'][0]['text']
                
                # Add translated text to new document
                translated_doc.add_paragraph(translated_text)
                
            except requests.exceptions.RequestException as e:
                print(f"Error during translation: {e}")
                return None
    
    # Save the translated document with "_traduzido" suffix
    output_path = os.path.splitext(documento)[0] + "_traduzido.docx"
    translated_doc.save(output_path)
    return output_path

if __name__ == "__main__":
    input_doc = input("Enter the path to the document to translate: ")
    if os.path.exists(input_doc):
        result = traduzir_documento(input_doc)
        if result:
            print(f"Translation completed. Saved as: {result}")
        else:
            print("Translation failed.")
    else:
        print("File not found.")
